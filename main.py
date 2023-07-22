import sys

from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QTableWidget, QTableWidgetItem, QVBoxLayout, QAction, \
    QFileDialog, QMenu, QPushButton, QHBoxLayout, QLabel, QMessageBox, QDialog, QRadioButton, QButtonGroup, \
    QDesktopWidget
from PyQt5.QtGui import QIcon
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


class LanguageSelectionDialog(QDialog):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Вибір мови")
        self.setGeometry(600, 300, 200, 100)
        self.language_label = QLabel("Оберіть мову:")
        self.ukrainian_radio = QRadioButton("Українська")
        self.english_radio = QRadioButton("English")
        self.ukrainian_radio.setChecked(True)

        self.button_group = QButtonGroup(self)
        self.button_group.addButton(self.ukrainian_radio, 1)
        self.button_group.addButton(self.english_radio, 2)

        self.confirm_button = QPushButton("Підтвердити")
        self.confirm_button.clicked.connect(self.accept)

        layout = QVBoxLayout()
        layout.addWidget(self.language_label)
        layout.addWidget(self.ukrainian_radio)
        layout.addWidget(self.english_radio)
        layout.addWidget(self.confirm_button)

        self.setLayout(layout)

    def selected_language(self):
        if self.button_group.checkedId() == 1:
            return "ua"
        else:
            return "en"

    def center(self):
        qr = self.frameGeometry()
        cp = QDesktopWidget().availableGeometry().center()
        qr.moveCenter(cp)


class TableEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.language_dialog = LanguageSelectionDialog()
        self.language_dialog.exec_()
        self.language = self.language_dialog.selected_language()
        self.init_ui()

    def init_ui(self):
        # Заголовок і розміри вікна
        if self.language == "ua":
            title = "Редактор таблиць"
        else:
            title = "Table Editor"
        self.setWindowTitle(title)
        self.setWindowState(Qt.WindowMaximized)
        # self.setGeometry(100, 100, 800, 600)

        # Створення таблиці 1х1
        self.table = QTableWidget(self)
        self.table.setRowCount(1)
        self.table.setColumnCount(1)

        # Створення кнопок для додавання/видалення рядків та стовпців
        save_button = QPushButton("Зберегти", self)
        open_button = QPushButton("Відкрити", self)
        add_row_button = QPushButton("Додати рядок", self)
        add_column_button = QPushButton("Додати стовпець", self)
        remove_row_button = QPushButton("Видалити рядок", self)
        remove_column_button = QPushButton("Видалити стовпець", self)

        if self.language == "ua":
            file_menu = self.menuBar().addMenu("Файл")
            save_action = QAction(QIcon(), "Зберегти", self)
            open_action = QAction(QIcon(), "Відкрити", self)
            exit_action = QAction(QIcon(), "Закрити", self)
            settings_menu = self.menuBar().addMenu("Налаштування")
            theme_menu = QMenu("Тема", self)
            dark_theme_action = QAction(QIcon(), "Темна", self)
            light_theme_action = QAction(QIcon(), "Світла", self)
            save_button.setText("Зберегти")
            open_button.setText("Відкрити")
            add_row_button.setText("Додати рядок")
            add_column_button.setText("Додати стовпець")
            remove_row_button.setText("Видалити рядок")
            remove_column_button.setText("Видалити стовпець")
            self.status_label = QLabel("Готово", self)
        else:
            file_menu = self.menuBar().addMenu("File")
            save_action = QAction(QIcon(), "Save", self)
            open_action = QAction(QIcon(), "Open", self)
            exit_action = QAction(QIcon(), "Exit", self)
            save_button.setText("Save")
            open_button.setText("Open")
            settings_menu = self.menuBar().addMenu("Settings")
            theme_menu = QMenu("Theme", self)
            dark_theme_action = QAction(QIcon(), "Dark", self)
            light_theme_action = QAction(QIcon(), "Light", self)
            add_row_button.setText("Add Row")
            add_column_button.setText("Add Column")
            remove_row_button.setText("Remove Row")
            remove_column_button.setText("Remove Column")
            self.status_label = QLabel("Ready", self)

        save_button.clicked.connect(self.save_file)
        open_button.clicked.connect(self.open_file)
        add_row_button.clicked.connect(self.add_row)
        add_column_button.clicked.connect(self.add_column)
        remove_row_button.clicked.connect(self.remove_row)
        remove_column_button.clicked.connect(self.remove_column)

        # Створення меню File
        save_action.triggered.connect(self.save_file)
        file_menu.addAction(save_action)

        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)

        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        dark_theme_action.triggered.connect(lambda: self.set_theme("dark"))
        theme_menu.addAction(dark_theme_action)

        light_theme_action.triggered.connect(lambda: self.set_theme("light"))
        theme_menu.addAction(light_theme_action)
        settings_menu.addMenu(theme_menu)

        # Групування кнопок додавання/видалення рядків та стовпців
        button_layout = QHBoxLayout()
        button_layout.addWidget(add_row_button)
        button_layout.addWidget(remove_row_button)
        button_layout.addWidget(add_column_button)
        button_layout.addWidget(remove_column_button)

        # Групування всіх компонентів
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.table)
        main_layout.addLayout(button_layout)
        # main_layout.addLayout(file_button_layout)
        main_layout.addWidget(self.status_label)

        # Основний віджет
        central_widget = QWidget(self)
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)

        # Тема за замовчуванням
        self.set_theme("light")

        # Мова за замовчуванням
        # self.set_language("en_US")

    def add_row(self):
        self.table.setRowCount(self.table.rowCount() + 1)

    def remove_row(self):
        if self.table.rowCount() > 1:
            self.table.setRowCount(self.table.rowCount() - 1)

    def add_column(self):
        self.table.setColumnCount(self.table.columnCount() + 1)

    def remove_column(self):
        if self.table.columnCount() > 1:
            self.table.setColumnCount(self.table.columnCount() - 1)

    def save_file(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Word Documents (*.docx);;All Files (*)",
                                                   options=options)
        if file_name:
            document = Document()
            table = document.add_table(rows=self.table.rowCount(), cols=self.table.columnCount())
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row in range(self.table.rowCount()):
                for col in range(self.table.columnCount()):
                    item = self.table.item(row, col)
                    if item:
                        table.cell(row, col).text = item.text()

            try:
                document.save(file_name)
                if self.language == "ua":
                    QMessageBox.information(self, "Успіх", "Файл успішно збережено.", QMessageBox.Ok)
                else:
                    QMessageBox.information(self, "Success", "File saved successfully.", QMessageBox.Ok)
            except Exception as e:
                if self.language == "ua":
                    QMessageBox.critical(self, "Помилка", f"При збереженні виникла помилка: {str(e)}", QMessageBox.Ok)
                else:
                    QMessageBox.critical(self, "Error", f"Error while saving file: {str(e)}", QMessageBox.Ok)

    def open_file(self):
        options = QFileDialog.Options()
        if self.language == "ua":
            open_file_lang_dialog = "Відкрити файл"
            open_file_lang_success = "Файл відкрито успішно."
            open_file_lang_error = "Помилка при відкритті файлу:"
            success_msg = "Успіх"
            error_msg = "Помилка"
        else:
            open_file_lang_dialog = "Open File"
            open_file_lang_success = "File opened successfully."
            open_file_lang_error = "Error while opening file:"
            success_msg = "Success"
            error_msg = "Error"
        file_name, _ = QFileDialog.getOpenFileName(self, open_file_lang_dialog, "",
                                                   "Word Documents (*.docx);;All Files (*)",
                                                   options=options)
        if file_name:
            try:
                document = Document(file_name)
                table = document.tables[0]
                self.table.setRowCount(len(table.rows))
                self.table.setColumnCount(len(table.columns))
                for row in range(len(table.rows)):
                    for col in range(len(table.columns)):
                        self.table.setItem(row, col, QTableWidgetItem(table.cell(row, col).text))
                self.status_label.setText(open_file_lang_success)
                QMessageBox.information(self, success_msg, open_file_lang_success, QMessageBox.Ok)
            except Exception as e:
                QMessageBox.critical(self, error_msg, f"{open_file_lang_error} {str(e)}", QMessageBox.Ok)

    def set_theme(self, theme):
        if theme == "dark":
            self.setStyleSheet("background-color: #333; color: #FFF;")
        else:
            self.setStyleSheet("background-color: #FFF; color: #000;")


# Запуск додатку
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = TableEditor()
    window.show()
    sys.exit(app.exec_())
